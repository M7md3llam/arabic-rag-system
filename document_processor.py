"""
Document processor for extracting text from various file formats
"""
import os
from pathlib import Path
from typing import Dict, List, Tuple
import PyPDF2
from docx import Document
import openpyxl
from PIL import Image
import io

class DocumentProcessor:
    """Process different document types and extract text"""
    
    def __init__(self, ocr_processor=None):
        self.ocr_processor = ocr_processor
        self.supported_types = {
            'application/pdf': self.process_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self.process_docx,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self.process_xlsx,
            'image/png': self.process_image,
            'image/jpeg': self.process_image,
            'image/jpg': self.process_image,
        }
    
    def process_file(self, file_path: str, file_type: str) -> Dict:
        """
        Process a file and extract text content
        
        Returns:
            Dict with keys: text, pages, metadata, status
        """
        try:
            processor = self.supported_types.get(file_type)
            if not processor:
                return {
                    'text': '',
                    'pages': [],
                    'metadata': {},
                    'status': 'unsupported_type',
                    'error': f'Unsupported file type: {file_type}'
                }
            
            result = processor(file_path)
            result['status'] = 'success'
            return result
            
        except Exception as e:
            return {
                'text': '',
                'pages': [],
                'metadata': {},
                'status': 'error',
                'error': str(e)
            }
    
    def process_pdf(self, file_path: str) -> Dict:
        """Extract text from PDF - with OCR fallback for scanned PDFs"""
        pages = []
        full_text = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                # Try extracting text first
                total_text_length = 0
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    total_text_length += len(text.strip())
                    
                    pages.append({
                        'page_number': page_num + 1,
                        'text': text,
                        'char_count': len(text)
                    })
                    full_text.append(text)
                
                # If very little text extracted, likely scanned PDF
                avg_chars_per_page = total_text_length / num_pages if num_pages > 0 else 0
                
                if avg_chars_per_page < 30 and self.ocr_processor:
                    # Use OCR for scanned PDF
                    print(f"Detected scanned PDF (avg {avg_chars_per_page:.0f} chars/page), using OCR...")
                    ocr_result = self.ocr_processor.process_scanned_pdf(file_path, max_pages=min(num_pages, 20))
                    
                    if ocr_result['status'] == 'success':
                        return {
                            'text': ocr_result['text'],
                            'pages': ocr_result['pages'],
                            'metadata': {
                                'num_pages': num_pages,
                                'type': 'pdf',
                                'method': 'ocr',
                                'ocr_pages': len(ocr_result['pages'])
                            }
                        }
        
        except Exception as e:
            print(f"Error in PDF processing: {e}")
        
        return {
            'text': '\n\n'.join(full_text),
            'pages': pages,
            'metadata': {
                'num_pages': len(pages),
                'type': 'pdf',
                'method': 'text_extraction'
            }
        }
    
    def process_docx(self, file_path: str) -> Dict:
        """Extract text from Word document"""
        doc = Document(file_path)
        
        paragraphs = []
        full_text = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append({
                    'text': para.text,
                    'style': para.style.name
                })
                full_text.append(para.text)
        
        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        return {
            'text': '\n\n'.join(full_text),
            'pages': paragraphs,
            'metadata': {
                'num_paragraphs': len(paragraphs),
                'num_tables': len(tables),
                'tables': tables,
                'type': 'docx'
            }
        }
    
    def process_xlsx(self, file_path: str) -> Dict:
        """Extract text from Excel file"""
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        
        sheets_data = []
        full_text = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            
            # Get all values
            sheet_text = []
            sheet_data = []
            
            for row in sheet.iter_rows(values_only=True):
                row_data = [str(cell) if cell is not None else '' for cell in row]
                sheet_data.append(row_data)
                row_text = ' | '.join([cell for cell in row_data if cell])
                if row_text:
                    sheet_text.append(row_text)
            
            sheets_data.append({
                'sheet_name': sheet_name,
                'text': '\n'.join(sheet_text),
                'data': sheet_data
            })
            full_text.extend(sheet_text)
        
        return {
            'text': '\n\n'.join(full_text),
            'pages': sheets_data,
            'metadata': {
                'num_sheets': len(sheets_data),
                'sheet_names': [s['sheet_name'] for s in sheets_data],
                'type': 'xlsx'
            }
        }
    
    def process_image(self, file_path: str) -> Dict:
        """
        Process image file with OCR
        """
        try:
            img = Image.open(file_path)
            
            # If OCR processor available, use it
            if self.ocr_processor:
                ocr_result = self.ocr_processor.process_image(file_path)
                
                if ocr_result['status'] == 'success':
                    return {
                        'text': ocr_result['text'],
                        'pages': [{
                            'type': 'image',
                            'size': img.size,
                            'format': img.format,
                            'mode': img.mode,
                            'ocr_text': ocr_result['text']
                        }],
                        'metadata': {
                            'width': img.width,
                            'height': img.height,
                            'format': img.format,
                            'type': 'image',
                            'method': 'gpt4_vision_ocr'
                        }
                    }
            
            # Fallback if no OCR
            return {
                'text': '[Image file - OCR not available]',
                'pages': [{
                    'type': 'image',
                    'size': img.size,
                    'format': img.format,
                    'mode': img.mode
                }],
                'metadata': {
                    'width': img.width,
                    'height': img.height,
                    'format': img.format,
                    'type': 'image'
                }
            }
        except Exception as e:
            return {
                'text': '',
                'pages': [],
                'metadata': {},
                'error': str(e)
            }
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """
        Split text into chunks with overlap
        
        Args:
            text: Input text
            chunk_size: Size of each chunk in characters
            overlap: Overlap between chunks
        
        Returns:
            List of text chunks
        """
        if not text:
            return []
        
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = start + chunk_size
            
            # Try to break at sentence boundary
            if end < text_length:
                # Look for sentence endings
                for punct in ['. ', '。', '！', '؟', '\n\n']:
                    last_punct = text[start:end].rfind(punct)
                    if last_punct != -1:
                        end = start + last_punct + len(punct)
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap
        
        return chunks


# Test function
if __name__ == "__main__":
    processor = DocumentProcessor()
    
    # Test with a sample file
    test_file = "sample.pdf"  # Replace with actual file path
    if os.path.exists(test_file):
        result = processor.process_file(test_file, 'application/pdf')
        print(f"Status: {result['status']}")
        print(f"Text length: {len(result['text'])}")
        print(f"Number of pages: {len(result['pages'])}")